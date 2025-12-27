import pdfplumber
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def convert_pdf_to_docx(pdf_path, docx_path):
    LINE_Y_TOLERANCE = 2
    FONT_NAME = "Times New Roman"
    FONT_SIZE = 12
    HEADING_SIZE = 14
    HEADING_KEYWORDS = ["FORM", "DETAILS OF PARTIES", "DETAILS OF DISPUTE"]

    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_NAME
    font.size = Pt(FONT_SIZE)

    pdf_lines = []

    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            words = page.extract_words()

            lines = {}
            for w in words:
                y_top = round(w['top'])
                added = False
                for key in lines:
                    if abs(key - y_top) <= LINE_Y_TOLERANCE:
                        lines[key].append(w)
                        added = True
                        break
                if not added:
                    lines[y_top] = [w]

            for y in sorted(lines.keys()):
                line_words = sorted(lines[y], key=lambda x: x['x0'])
                line_text = ""
                for w in line_words:
                    if line_text.endswith("{") or w['text'].startswith("}") or w['text'].startswith("%"):
                        line_text += w['text']
                    else:
                        line_text += (" " + w['text']) if line_text else w['text']
                pdf_lines.append(line_text)

    # Merge numbering
    merged_lines = []
    i = 0
    while i < len(pdf_lines):
        line = pdf_lines[i].strip()
        if line.isdigit() and i + 1 < len(pdf_lines):
            merged_lines.append(line + " " + pdf_lines[i + 1].strip())
            i += 2
        else:
            merged_lines.append(line)
            i += 1

    # Merge placeholders
    final_lines = []
    i = 0
    while i < len(merged_lines):
        line = merged_lines[i]
        if line.strip().startswith("{%") and not line.strip().endswith("%}"):
            block = line
            i += 1
            while i < len(merged_lines) and not merged_lines[i].strip().endswith("%}"):
                block += " " + merged_lines[i].strip()
                i += 1
            if i < len(merged_lines):
                block += " " + merged_lines[i].strip()
                i += 1
            final_lines.append(block)
        else:
            final_lines.append(line)
            i += 1

    # Write to Word
    for line in final_lines:
        if any(keyword in line.upper() for keyword in HEADING_KEYWORDS):
            p = doc.add_paragraph(line)
            run = p.runs[0]
            run.bold = True
            run.font.size = Pt(HEADING_SIZE)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        else:
            p = doc.add_paragraph(line)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    doc.save(docx_path)
